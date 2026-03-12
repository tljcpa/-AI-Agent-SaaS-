"""Office Graph 适配器：操作 OneDrive 中的 Word/Excel。"""
from __future__ import annotations

import string
from io import BytesIO

import httpx
from docx import Document

from app.adapters.ms_auth import MSAuthService


def _col_letter(n: int) -> str:
    result = ""
    while n:
        n, remainder = divmod(n - 1, 26)
        result = string.ascii_uppercase[remainder] + result
    return result


class GraphOfficeProvider:
    def __init__(self, auth_service: MSAuthService, http_client: httpx.AsyncClient, root_path: str = "") -> None:
        self.auth_service = auth_service
        self.http_client = http_client
        self.root_path = root_path.strip("/")

    async def _headers(self, user_id: int) -> dict[str, str]:
        token = await self.auth_service.get_valid_access_token(user_id)
        return {"Authorization": f"Bearer {token}"}

    def _content_url(self, user_id: int, file_id: str) -> str:
        """拼接文件内容访问 URL（路径风格，兼容文件名作为 file_id）。"""
        prefix = f"{self.root_path}/" if self.root_path else ""
        path = f"{prefix}user_{user_id}/{file_id}"
        return f"https://graph.microsoft.com/v1.0/me/drive/root:/{path}:/content"

    async def _resolve_item_id(self, user_id: int, file_id: str) -> str:
        """通过路径解析出 OneDrive item ID，用于 workbook API。"""
        headers = await self._headers(user_id)
        prefix = f"{self.root_path}/" if self.root_path else ""
        path = f"{prefix}user_{user_id}/{file_id}"
        url = f"https://graph.microsoft.com/v1.0/me/drive/root:/{path}"
        resp = await self.http_client.get(url, headers=headers)
        resp.raise_for_status()
        return str(resp.json()["id"])

    async def read_word_content(self, user_id: int, file_id: str) -> str:
        headers = await self._headers(user_id)
        url = self._content_url(user_id, file_id)
        resp = await self.http_client.get(url, headers=headers)
        resp.raise_for_status()
        doc = Document(BytesIO(resp.content))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    async def format_word_document(self, user_id: int, file_id: str, style_instructions: str) -> str:
        headers = await self._headers(user_id)
        url = self._content_url(user_id, file_id)
        download = await self.http_client.get(url, headers=headers)
        download.raise_for_status()

        doc = Document(BytesIO(download.content))
        for para in doc.paragraphs:
            if para.text.strip().startswith("#"):
                para.style = "Heading 1"
            elif "标题" in style_instructions:
                para.style = "Heading 2" if len(para.text) < 24 else para.style
            for run in para.runs:
                if "字体" in style_instructions:
                    run.font.name = "Calibri"

        out = BytesIO()
        doc.save(out)
        out.seek(0)

        upload = await self.http_client.put(url, headers=headers, content=out.read())
        upload.raise_for_status()
        return f"Word 文档 {file_id} 已完成样式更新"

    async def read_excel_data(self, user_id: int, file_id: str, sheet_name: str) -> str:
        headers = await self._headers(user_id)
        item_id = await self._resolve_item_id(user_id, file_id)
        url = (
            "https://graph.microsoft.com/v1.0/me/drive/items/"
            f"{item_id}/workbook/worksheets/{sheet_name}/usedRange"
        )
        resp = await self.http_client.get(url, headers=headers)
        resp.raise_for_status()
        payload = resp.json()
        values = payload.get("values", [])
        return "\n".join([", ".join(map(str, row)) for row in values])

    async def write_excel_data(self, user_id: int, file_id: str, sheet_name: str, data: list[list]) -> str:
        if not data:
            return "Excel 数据为空，跳过写入"

        headers = await self._headers(user_id)
        item_id = await self._resolve_item_id(user_id, file_id)
        range_url = (
            "https://graph.microsoft.com/v1.0/me/drive/items/"
            f"{item_id}/workbook/worksheets/{sheet_name}/range(address='A1')"
        )
        chart_url = (
            "https://graph.microsoft.com/v1.0/me/drive/items/"
            f"{item_id}/workbook/worksheets/{sheet_name}/charts/add"
        )
        num_cols = len(data[0])
        end_col = _col_letter(num_cols)
        source_data = f"{sheet_name}!A1:{end_col}{max(2, len(data))}"

        range_resp = await self.http_client.patch(
            range_url,
            headers=headers,
            json={"values": data},
        )
        range_resp.raise_for_status()
        chart_resp = await self.http_client.post(
            chart_url,
            headers=headers,
            json={"type": "ColumnClustered", "sourceData": source_data, "seriesBy": "Auto"},
        )
        if chart_resp.status_code >= 400 and chart_resp.status_code != 409:
            chart_resp.raise_for_status()
        return f"Excel {sheet_name} 写入 {len(data)} 行并尝试更新图表"

