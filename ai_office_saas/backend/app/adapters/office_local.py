from __future__ import annotations

from pathlib import Path

import openpyxl
from docx import Document

from app.adapters.storage_local import LocalStorageProvider


class LocalOfficeProvider:
    def __init__(self, storage: LocalStorageProvider) -> None:
        self.storage = storage

    def _file_path(self, user_id: int, file_id: str) -> Path:
        user_root = self.storage._user_root(user_id)
        target = user_root / file_id
        self.storage._assert_in_sandbox(user_root, target)
        if not target.exists():
            raise FileNotFoundError(f"文件不存在: {file_id}")
        return target

    async def read_word_content(self, user_id: int, file_id: str) -> str:
        path = self._file_path(user_id, file_id)
        doc = Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        text = "\n".join(paragraphs)
        # 限制返回内容最多 8000 字符，避免超出 LLM context window
        if len(text) > 8000:
            text = text[:8000] + "\n...(内容过长，已截断)"
        return text if text else "（文档内容为空）"

    async def format_word_document(self, user_id: int, file_id: str, style_instructions: str) -> str:
        path = self._file_path(user_id, file_id)
        doc = Document(str(path))
        instructions_lower = style_instructions.lower()
        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            if para.text.startswith("#"):
                para.style = "Heading 1"
            for run in para.runs:
                if "加粗" in instructions_lower or "bold" in instructions_lower:
                    run.bold = True
                if "字体" in instructions_lower or "font" in instructions_lower:
                    run.font.name = "Calibri"
        doc.save(str(path))
        return f"Word 文档 {file_id} 已完成格式化"

    async def read_excel_data(self, user_id: int, file_id: str, sheet_name: str) -> str:
        path = self._file_path(user_id, file_id)
        wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
        if sheet_name not in wb.sheetnames:
            available = ", ".join(wb.sheetnames)
            wb.close()
            return f"Sheet '{sheet_name}' 不存在，可用 Sheet：{available}"
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            if any(cell is not None for cell in row):
                rows.append(", ".join("" if v is None else str(v) for v in row))
        wb.close()
        text = "\n".join(rows)
        if len(text) > 8000:
            text = text[:8000] + "\n...(内容过长，已截断)"
        return text if text else "（Sheet 内容为空）"

    async def write_excel_data(self, user_id: int, file_id: str, sheet_name: str, data: list) -> str:
        if not data:
            return "数据为空，跳过写入"
        path = self._file_path(user_id, file_id)
        if path.exists():
            wb = openpyxl.load_workbook(str(path))
        else:
            wb = openpyxl.Workbook()
            if wb.active:
                wb.active.title = sheet_name
        if sheet_name not in wb.sheetnames:
            wb.create_sheet(sheet_name)
        ws = wb[sheet_name]
        ws.delete_rows(1, ws.max_row)
        for row in data:
            ws.append(row if isinstance(row, list) else [str(row)])
        wb.save(str(path))
        return f"Excel {file_id} 的 Sheet '{sheet_name}' 已写入 {len(data)} 行"
